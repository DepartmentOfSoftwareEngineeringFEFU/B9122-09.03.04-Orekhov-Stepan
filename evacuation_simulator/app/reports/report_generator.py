from __future__ import annotations

from pathlib import Path

from docx import Document

from app.models import EvacuationScenario, SimulationResult


class ReportGenerator:

    def build_text(self, scenario: EvacuationScenario, result: SimulationResult | None) -> str:
        if result is None:
            raise ValueError("Результаты моделирования отсутствуют")
        lines: list[str] = []
        lines.append("Отчёт по результатам моделирования эвакуации")
        lines.append("")
        lines.append("1. Сведения о модели помещения")
        lines.append(f"Модель: {scenario.building_model.name}")
        lines.append(f"Размеры: {scenario.building_model.width} x {scenario.building_model.height} м")
        lines.append(f"Шаг клеточного разбиения: {scenario.building_model.cell_size} м")
        lines.append(f"Количество клеток: {scenario.building_model.cols * scenario.building_model.rows}")
        lines.append("")
        lines.append("2. Исходные параметры сценария")
        params = scenario.simulation_parameters
        lines.append(f"Сценарий: {scenario.name}")
        lines.append(f"Δt: {params.time_step} с; Tmax: {params.max_time} с")
        lines.append(f"alpha={params.alpha}, beta={params.beta}, gamma={params.gamma}, rho_crit={params.rho_crit}")
        lines.append("")
        lines.append("3. Параметры агентов")
        for agent in sorted(scenario.agents, key=lambda item: item.id):
            lines.append(
                f"Агент {agent.id}: группа {agent.mobility_group.value}, состояние {agent.state.value}, "
                f"f_eff={agent.effective_projection_area:.3f} м²/чел."
            )
        lines.append("")
        lines.append("4. Список выходов и их статусы")
        for exit_obj in sorted(scenario.exits, key=lambda item: item.id):
            lines.append(
                f"Выход {exit_obj.id}: ширина {exit_obj.width:.2f} м, статус {exit_obj.status.value}, "
                f"capacity={exit_obj.capacity}"
            )
        lines.append("")
        lines.append("5. События среды")
        if scenario.events:
            for event in scenario.events:
                lines.append(f"Событие {event.id}: {event.type.value}, t={event.time}")
        else:
            lines.append("События не заданы")
        lines.append("")
        lines.append("6. Итоговые результаты")
        tev = "не определено" if result.total_evacuation_time is None else f"{result.total_evacuation_time:.1f} с"
        lines.append(f"Tevac: {tev}")
        lines.append(f"Эвакуировано: {result.evacuated_count}")
        lines.append(f"Не эвакуировано: {result.blocked_count}")
        lines.append(f"Задержано: {result.delayed_count}")
        lines.append(f"Требуют спасения: {result.needs_rescue_count}")
        lines.append("")
        lines.append("7. Таблица времени эвакуации агентов")
        for agent in sorted(scenario.agents, key=lambda item: item.id):
            value = result.agent_evacuation_times.get(agent.id)
            text = "не эвакуирован" if value is None else f"{value:.1f} с"
            lines.append(f"Агент {agent.id}: {text}")
        lines.append("")
        lines.append("8. Описание зон заторов")
        lines.append("Зоны заторов:")
        self._append_zones(lines, result.congestion_zones)
        lines.append("")
        lines.append("9. Нормативные скопления Dnorm > 0.5 м²/м²")
        self._append_zones(lines, result.normative_congestion_zones)
        lines.append("")
        lines.append("10. Предупреждения")
        if result.warnings:
            lines.extend(result.warnings)
        else:
            lines.append("Предупреждения отсутствуют")
        lines.append("")
        lines.append("11. Вывод по сценарию")
        if params.evacuation_start_time is None or params.route_blocking_time is None:
            lines.append(
                "Сформированы результаты моделирования эвакуации без полной нормативной оценки пожарной безопасности, "
                "так как не заданы время начала эвакуации и/или время блокирования путей эвакуации."
            )
        elif result.blocked_count == 0:
            lines.append("Все самостоятельно движущиеся агенты эвакуированы в рамках заданного сценария.")
        else:
            lines.append("В сценарии имеются неэвакуированные или задержанные агенты.")
        if result.needs_rescue_count:
            lines.append(
                "Итоговый нормативный вывод является неполным, так как для групп НМ и НТ требуется отдельный расчёт спасения."
            )
        return "\n".join(lines)

    def save_docx(self, scenario: EvacuationScenario, result: SimulationResult | None, path: str | Path) -> None:
        if result is None:
            raise ValueError("Результаты моделирования отсутствуют")
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading("Отчёт по результатам моделирования эвакуации", level=1)
        self._add_model_section(doc, scenario)
        self._add_scenario_section(doc, scenario)
        self._add_agents_section(doc, scenario)
        self._add_exits_section(doc, scenario)
        self._add_events_section(doc, scenario)
        self._add_results_section(doc, result)
        self._add_evacuation_times_section(doc, scenario, result)
        self._add_zones_section(doc, "8. Описание зон заторов", result.congestion_zones)
        self._add_zones_section(doc, "9. Нормативные скопления Dnorm > 0.5 м²/м²", result.normative_congestion_zones)
        self._add_warnings_section(doc, result)
        self._add_conclusion_section(doc, scenario, result)
        doc.save(target)

    def _append_zones(self, lines: list[str], zones: list[dict[str, object]]) -> None:
        if not zones:
            lines.append("Не зафиксированы")
            return
        for zone in zones[:30]:
            lines.append(f"t={zone.get('time')} с, клетка={zone.get('coord')}, Dnorm={float(zone.get('dnorm', 0)):.3f}")
        if len(zones) > 30:
            lines.append(f"... ещё {len(zones) - 30} записей")

    def _add_model_section(self, doc: Document, scenario: EvacuationScenario) -> None:
        building = scenario.building_model
        doc.add_heading("1. Сведения о модели помещения", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Параметр"
        table.rows[0].cells[1].text = "Значение"
        for name, value in [
            ("Модель", building.name),
            ("Размеры помещения", f"{building.width} x {building.height} м"),
            ("Шаг клеточного разбиения", f"{building.cell_size} м"),
            ("Количество клеток", str(building.cols * building.rows)),
        ]:
            self._add_table_row(table, [name, value])

    def _add_scenario_section(self, doc: Document, scenario: EvacuationScenario) -> None:
        params = scenario.simulation_parameters
        doc.add_heading("2. Исходные параметры сценария", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Параметр"
        table.rows[0].cells[1].text = "Значение"
        for name, value in [
            ("Сценарий", scenario.name),
            ("Шаг моделирования Δt", f"{params.time_step} с"),
            ("Максимальное время Tmax", f"{params.max_time} с"),
            ("alpha", str(params.alpha)),
            ("beta", str(params.beta)),
            ("gamma", str(params.gamma)),
            ("rho_crit", str(params.rho_crit)),
        ]:
            self._add_table_row(table, [name, value])

    def _add_agents_section(self, doc: Document, scenario: EvacuationScenario) -> None:
        doc.add_heading("3. Параметры агентов", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        for index, header in enumerate(["ID", "Группа", "Одежда", "Скорость, м/с", "f_eff, м²/чел.", "Стартовая клетка"]):
            table.rows[0].cells[index].text = header
        for agent in sorted(scenario.agents, key=lambda item: item.id):
            self._add_table_row(
                table,
                [
                    agent.id,
                    agent.mobility_group.value,
                    agent.clothes_type.value,
                    f"{agent.base_speed:.2f}",
                    f"{agent.effective_projection_area:.3f}",
                    str(agent.start_cell),
                ],
            )

    def _add_exits_section(self, doc: Document, scenario: EvacuationScenario) -> None:
        doc.add_heading("4. Список выходов и их статусы", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for index, header in enumerate(["ID", "Клетки", "Ширина, м", "Статус", "Пропускная способность"]):
            table.rows[0].cells[index].text = header
        for exit_obj in sorted(scenario.exits, key=lambda item: item.id):
            self._add_table_row(
                table,
                [exit_obj.id, str(sorted(exit_obj.cells)), f"{exit_obj.width:.2f}", exit_obj.status.value, exit_obj.capacity],
            )

    def _add_events_section(self, doc: Document, scenario: EvacuationScenario) -> None:
        doc.add_heading("5. События среды", level=2)
        if not scenario.events:
            doc.add_paragraph("События не заданы.")
            return
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for index, header in enumerate(["ID", "Тип", "Время, с", "Параметры"]):
            table.rows[0].cells[index].text = header
        for event in sorted(scenario.events, key=lambda item: item.id):
            self._add_table_row(table, [event.id, event.type.value, event.time, event.params])

    def _add_results_section(self, doc: Document, result: SimulationResult) -> None:
        doc.add_heading("6. Итоговые результаты", level=2)
        tev = "не определено" if result.total_evacuation_time is None else f"{result.total_evacuation_time:.1f} с"
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Показатель"
        table.rows[0].cells[1].text = "Значение"
        for name, value in [
            ("Tevac", tev),
            ("Эвакуировано", result.evacuated_count),
            ("Не эвакуировано", result.blocked_count),
            ("Задержано", result.delayed_count),
            ("Требуют спасения", result.needs_rescue_count),
        ]:
            self._add_table_row(table, [name, value])

    def _add_evacuation_times_section(self, doc: Document, scenario: EvacuationScenario, result: SimulationResult) -> None:
        doc.add_heading("7. Таблица времени эвакуации агентов", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for index, header in enumerate(["ID", "Группа", "Состояние", "Время эвакуации"]):
            table.rows[0].cells[index].text = header
        for agent in sorted(scenario.agents, key=lambda item: item.id):
            value = result.agent_evacuation_times.get(agent.id)
            self._add_table_row(
                table,
                [agent.id, agent.mobility_group.value, agent.state.value, "не эвакуирован" if value is None else f"{value:.1f} с"],
            )

    def _add_zones_section(self, doc: Document, heading: str, zones: list[dict[str, object]]) -> None:
        doc.add_heading(heading, level=2)
        if not zones:
            doc.add_paragraph("Не зафиксированы.")
            return
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for index, header in enumerate(["Время, с", "Клетка", "Dnorm"]):
            table.rows[0].cells[index].text = header
        for zone in zones[:30]:
            self._add_table_row(table, [zone.get("time"), zone.get("coord"), f"{float(zone.get('dnorm', 0)):.3f}"])
        if len(zones) > 30:
            doc.add_paragraph(f"Показаны первые 30 записей из {len(zones)}.")

    def _add_warnings_section(self, doc: Document, result: SimulationResult) -> None:
        doc.add_heading("10. Предупреждения", level=2)
        if result.warnings:
            for warning in result.warnings:
                doc.add_paragraph(warning, style="List Bullet")
        else:
            doc.add_paragraph("Предупреждения отсутствуют.")

    def _add_conclusion_section(self, doc: Document, scenario: EvacuationScenario, result: SimulationResult) -> None:
        doc.add_heading("11. Вывод по сценарию", level=2)
        params = scenario.simulation_parameters
        if params.evacuation_start_time is None or params.route_blocking_time is None:
            doc.add_paragraph(
                "Сформированы результаты моделирования эвакуации без полной нормативной оценки пожарной безопасности, "
                "так как не заданы время начала эвакуации и/или время блокирования путей эвакуации."
            )
        elif result.blocked_count == 0:
            doc.add_paragraph("Все самостоятельно движущиеся агенты эвакуированы в рамках заданного сценария.")
        else:
            doc.add_paragraph("В сценарии имеются неэвакуированные или задержанные агенты.")
        if result.needs_rescue_count:
            doc.add_paragraph(
                "Итоговый нормативный вывод является неполным, так как для групп НМ и НТ требуется отдельный расчёт спасения."
            )

    def _add_table_row(self, table, values: list[object]) -> None:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
